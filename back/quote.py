import pythoncom
from flask import Flask, request, send_file, jsonify
from docx import Document
import os
import json
import logging
from google.cloud import aiplatform
from google.auth import default
from docx2pdf import convert
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor, Pt
from docx.enum.style import WD_STYLE_TYPE
import vertexai
from vertexai.preview.generative_models import GenerativeModel, ChatSession
from flask_cors import CORS

# Set up logging
logging.basicConfig(level=logging.DEBUG)

app = Flask(__name__)
CORS(app)

# Set the project ID and service account key file
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "./intelligent-quoter-46583564dfe3.json"
project_id = "intelligent-quoter"
location = "us-central1"

# Initialize the Vertex AI client
vertexai.init(project=project_id, location=location)

model_name = "gemini-1.5-pro-001"

@app.route('/generate-quote', methods=['POST'])
def generate_quote():
    try:
        pythoncom.CoInitialize()
        user_request = request.json.get("request")
        logging.debug(f"User request: {user_request}")
        
        # Define the messages for Vertex AI
        messages = [
            {
                "role": "system",
                "content": (
                    "You are an expert AI that generates detailed quotes for web development projects in JSON format. "
                    "The JSON object must follow this structure: "
                    "{\n"
                    "  \"title\": \"string\",\n"
                    "  \"analysis\": \"string\",\n"
                    "  \"numbered_list\": [\n"
                    "    {\n"
                    "      \"page\": \"string\",\n"
                    "      \"content\": \"string\",\n"
                    "      \"subpages\": [\n"
                    "        {\"page\": \"string\", \"content\": \"string\"},\n"
                    "        ...\n"
                    "      ]\n"
                    "    },\n"
                    "    ...\n"
                    "  ],\n"
                    "  \"requirements\": [\"string\", ...],\n"
                    "  \"payment_details\": [\n"
                    "    {\"type\": \"string\", \"description\": \"string\", \"price\": \"string\"},\n"
                    "    ...\n"
                    "  ],\n"
                    "  \"final_price\": \"string\",\n"
                    "  \"project_duration\": \"string\"\n"
                    "}\n"
                    "Here are the fields:\n"
                    "- title: The title of the quote\n"
                    "- analysis: A detailed analysis of the project, including problem description and project solution (200-300 words)\n"
                    "- numbered_list: A list of pages and subpages. Main pages are numbered 1., 2., etc. Subpages, if any, are numbered 1.1, 1.2, etc.\n"
                    "- requirements: Dot points of what is required of the client to complete the project (shouldn't be much and should be concise)\n"
                    "- payment_details: A list of items charged, detailed into 3 fields respectively: Included or Optional, description of work to be completed, price in $\n"
                    "- final_price: The total price of the project for all included items\n"
                    "- project_duration: The estimated duration to complete the project\n"
                    "Make sure the JSON is valid and matches this structure."
                )
            },
            {
                "role": "user",
                "content": user_request
            }
        ]

        # Log the messages sent to Vertex AI
        logging.debug(f"Messages to Vertex AI: {messages}")

        # Generate content using Vertex AI
        model = GenerativeModel(model_name)
        chat = model.start_chat()
        prompt = json.dumps(messages)
        response = chat.send_message(prompt)

        # Log the raw response from Vertex AI
        logging.debug(f"Raw response from Vertex AI: {response.text}")

        # Clean the response to extract JSON part
        cleaned_response = response.text.strip().strip("```json").strip("```").strip()
        logging.debug(f"Cleaned response: {cleaned_response}")

        content = json.loads(cleaned_response)

        # Log the parsed JSON content
        logging.debug(f"Parsed content: {json.dumps(content, indent=2)}")

        # Add final price to payment details
        content['payment_details'].append({
            "type": "TOTAL",
            "description": "Without Options",
            "price": content['final_price']
        })

        # Load the Word template
        doc = Document('quote_template.docx')

        # Replace placeholders with actual content
        replace_placeholders(doc, content)

        # Save the generated document
        output_path = 'output_quote.docx'
        doc.save(output_path)
        
        # Convert the docx to pdf
        pdf_path = 'output_quote.pdf'
        convert(output_path, pdf_path)

        # Return the generated PDF document
        return send_file(pdf_path, as_attachment=True)

    except Exception as e:
        logging.error(f"Error: {str(e)}", exc_info=True)
        return jsonify({"error": str(e)}), 500
    finally:
        pythoncom.CoUninitialize()

def get_payment_plan(final_price):
    try:
        price = float(final_price.replace('$', '').replace(',', ''))
    except ValueError:
        return "Invalid price format"

    if price < 5000:
        return (
            "1. Start Project – 30% of Total Cost Transferred\n"
            "2. First Draft Approved – 40% of Total Cost Transferred\n"
            "3. Project Finalised – 30% of Total Cost Transferred"
        )
    elif price < 10000:
        return (
            "1. Start Project – 20% of Total Cost transferred\n"
            "2. First Draft Approved – 40% of Total Cost Transferred\n"
            "3. All Finishing Touches Approved and final version ready for deployment – 20% of Total Cost Transferred\n"
            "4. Project Finalised – 20% of Total Cost Transferred"
        )
    else:
        return (
            "1. Start Project - 10% of Total Cost Transferred\n"
            "2. Designs and Wireframes Approved - 20% of Total Cost Transferred\n"
            "3. First Draft Approved - 30% of Total Cost Transferred\n"
            "4. Second Draft Approved 10% of Total Cost Transferred\n"
            "5. All Finishing Touches Approved and final version ready for deployment – 15% of Total Cost Transferred\n"
            "6. Project Finalised – 15% of Total Cost Transferred"
        )

def replace_placeholders(doc, content):
    def replace_in_paragraph(paragraph, replacements):
        for key, value in replacements.items():
            placeholder = f'{{{{ {key} }}}}'
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    def replace_in_table(table, replacements):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)

    def handle_numbered_list(doc, numbered_list):
        # Create or get the styles we need
        styles = doc.styles
        main_list_style = styles.add_style('CustomMainNumberedList', WD_STYLE_TYPE.PARAGRAPH)
        main_list_style.base_style = styles['Normal']
        main_list_style.font.size = Pt(11)  # Adjust size as needed

        sub_list_style = styles.add_style('CustomSubNumberedList', WD_STYLE_TYPE.PARAGRAPH)
        sub_list_style.base_style = styles['Normal']
        sub_list_style.font.size = Pt(11)  # Adjust size as needed

        # Find the paragraph with the placeholder and replace it with our formatted list
        for paragraph in doc.paragraphs:
            if '{{ numbered_list }}' in paragraph.text:
                # Clear the placeholder text
                paragraph.text = ""
                
                # Add each item in the numbered list
                for item in numbered_list:
                    # Add main page
                    new_para = paragraph.insert_paragraph_before(f"{item['page']}. {item['content']}", style='CustomMainNumberedList')
                    new_para.style.paragraph_format.left_indent = Pt(0)
                    new_para.style.paragraph_format.first_line_indent = Pt(18)

                    # Add subpages if they exist
                    if 'subpages' in item and item['subpages']:
                        for subitem in item['subpages']:
                            sub_para = paragraph.insert_paragraph_before(f"{subitem['page']} {subitem['content']}", style='CustomSubNumberedList')
                            sub_para.style.paragraph_format.left_indent = Pt(36)
                            sub_para.style.paragraph_format.first_line_indent = Pt(18)

                # Remove the original placeholder paragraph
                p = paragraph._element
                p.getparent().remove(p)
                break

    def handle_requirements(doc, requirements):
        requirements_text = "\n".join([f"- {req}" for req in requirements])
        for paragraph in doc.paragraphs:
            if '{{ requirements }}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{ requirements }}', requirements_text)

    def handle_payment_details(doc, payment_details):
        for table in doc.tables:
            # Find the table with the payment details placeholders
            start_row = None
            for i, row in enumerate(table.rows):
                if any('{{ type_0 }}' in cell.text for cell in row.cells):
                    start_row = i
                    break

            if start_row is not None:
                # Remove existing rows from start_row+1 to the end
                for _ in range(len(table.rows) - start_row - 1):
                    table._element.remove(table.rows[-1]._element)

                # Update the start row with the first payment detail
                for j, cell in enumerate(table.rows[start_row].cells):
                    cell.text = payment_details[0][['type', 'description', 'price'][j]]
                    shading_elm = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="FFFFFF"/>')
                    cell._element.get_or_add_tcPr().append(shading_elm)

                # Add rows for remaining payment details with alternating colors
                for i, item in enumerate(payment_details[1:-1], start=1):  # Exclude the first and last items
                    new_row = table.add_row()
                    for j, cell in enumerate(new_row.cells):
                        cell.text = item[['type', 'description', 'price'][j]]
                        # Set background color
                        shading_elm = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{["FFFFFF", "E6F3F7"][i % 2]}"/>')
                        cell._element.get_or_add_tcPr().append(shading_elm)

                # Add the final black row for TOTAL
                total_row = table.add_row()
                for j, cell in enumerate(total_row.cells):
                    cell.text = payment_details[-1][['type', 'description', 'price'][j]]
                    # Set background color to black and text color to white
                    shading_elm = parse_xml('<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="000000"/>')
                    cell._element.get_or_add_tcPr().append(shading_elm)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)

    # Replace simple fields
    simple_fields = ['title', 'analysis', 'final_price', 'project_duration']
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, {key: content[key] for key in simple_fields if key in content})

    # Handle complex fields
    handle_numbered_list(doc, content['numbered_list'])
    handle_requirements(doc, content['requirements'])
    handle_payment_details(doc, content['payment_details'])

    # Replace payment plan
    payment_plan = get_payment_plan(content['final_price'])
    for paragraph in doc.paragraphs:
        if '{{ payment_plan }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ payment_plan }}', payment_plan)

    # Replace in tables
    for table in doc.tables:
        replace_in_table(table, content)
        
if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5001)